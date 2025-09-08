"""
Template Analyzer - Extract styling from Word documents
"""

import json
import os
from docx import Document
from docx.shared import Pt, Inches
from typing import Dict, Any, Optional
import logging

logger = logging.getLogger(__name__)


class TemplateAnalyzer:
    """Extract comprehensive styling information from Word documents"""

    def __init__(self):
        self.template_stats = {}

    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """Extract all styling information from a Word document"""
        doc = Document(file_path)

        stats = {
            'document_info': self._extract_document_info(doc),
            'sections': self._extract_sections(doc),
            'styles': self._extract_styles(doc),
            'paragraphs': self._extract_paragraph_samples(doc),
            'formatting_patterns': self._identify_patterns(doc)
        }

        return stats

    def _extract_document_info(self, doc: Document) -> Dict:
        """Extract basic document information"""
        return {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'core_properties': {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or ''
            }
        }

    def _extract_sections(self, doc: Document) -> Dict:
        """Extract page setup and margin information"""
        section = doc.sections[0]

        return {
            'page_width': float(section.page_width.inches),
            'page_height': float(section.page_height.inches),
            'margins': {
                'top': float(section.top_margin.inches),
                'bottom': float(section.bottom_margin.inches),
                'left': float(section.left_margin.inches),
                'right': float(section.right_margin.inches)
            }
        }

    def _extract_styles(self, doc: Document) -> Dict:
        """Extract all document styles"""
        styles = {}

        for style in doc.styles:
            style_info = {
                'name': style.name,
                'type': str(style.type),
                'builtin': style.builtin
            }

            # Extract font information
            if hasattr(style, 'font') and style.font:
                style_info['font'] = {
                    'name': style.font.name,
                    'size': float(style.font.size.pt) if style.font.size else None,
                    'bold': style.font.bold,
                    'italic': style.font.italic,
                    'color': str(style.font.color.rgb) if style.font.color and style.font.color.rgb else None
                }

            # Extract paragraph formatting
            if hasattr(style, 'paragraph_format') and style.paragraph_format:
                pf = style.paragraph_format
                style_info['paragraph_format'] = {
                    'alignment': str(pf.alignment) if pf.alignment else None,
                    'left_indent': float(pf.left_indent.inches) if pf.left_indent else None,
                    'first_line_indent': float(pf.first_line_indent.inches) if pf.first_line_indent else None,
                    'space_before': float(pf.space_before.pt) if pf.space_before else None,
                    'space_after': float(pf.space_after.pt) if pf.space_after else None,
                    'line_spacing': pf.line_spacing
                }

            styles[style.name] = style_info

        return styles

    def _extract_paragraph_samples(self, doc: Document) -> list:
        """Extract formatting from first 20 paragraphs as samples"""
        samples = []

        for i, paragraph in enumerate(doc.paragraphs[:20]):
            if not paragraph.text.strip():
                continue

            sample = {
                'index': i,
                'text_preview': paragraph.text[:50],
                'style_name': paragraph.style.name,
                'alignment': str(paragraph.alignment) if paragraph.alignment else None,
                'runs': []
            }

            # Extract run formatting
            for run in paragraph.runs[:3]:  # First 3 runs
                run_info = {
                    'text': run.text[:20],
                    'font_name': run.font.name,
                    'font_size': float(run.font.size.pt) if run.font.size else None,
                    'bold': run.font.bold,
                    'italic': run.font.italic
                }
                sample['runs'].append(run_info)

            samples.append(sample)

        return samples

    def _identify_patterns(self, doc: Document) -> Dict:
        """Identify common formatting patterns"""
        patterns = {
            'heading_styles': [],
            'body_styles': [],
            'list_styles': [],
            'common_fonts': {},
            'common_sizes': {}
        }

        # Analyze paragraph styles
        for para in doc.paragraphs[:50]:
            style_name = para.style.name

            if 'heading' in style_name.lower():
                patterns['heading_styles'].append(style_name)
            elif 'list' in style_name.lower():
                patterns['list_styles'].append(style_name)
            else:
                patterns['body_styles'].append(style_name)

            # Track fonts and sizes
            for run in para.runs:
                if run.font.name:
                    patterns['common_fonts'][run.font.name] = patterns['common_fonts'].get(run.font.name, 0) + 1
                if run.font.size:
                    size = float(run.font.size.pt)
                    patterns['common_sizes'][size] = patterns['common_sizes'].get(size, 0) + 1

        # Remove duplicates and sort by frequency
        patterns['heading_styles'] = list(set(patterns['heading_styles']))
        patterns['body_styles'] = list(set(patterns['body_styles']))
        patterns['list_styles'] = list(set(patterns['list_styles']))

        return patterns

    def save_analysis(self, stats: Dict, output_file: str):
        """Save analysis to JSON file"""
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(stats, f, indent=2, ensure_ascii=False, default=str)
            logger.info(f"Template analysis saved to {output_file}")
        except Exception as e:
            logger.error(f"Error saving analysis: {str(e)}")
            raise


# Usage
if __name__ == "__main__":
    analyzer = TemplateAnalyzer()

    # Analyze template
    template_path = "data/template.docx"  # Your template file
    if os.path.exists(template_path):
        stats = analyzer.analyze_document(template_path)
        analyzer.save_analysis(stats, "template_analysis.json")
        print("✅ Template analyzed and saved to template_analysis.json")
    else:
        print(f"❌ Template file not found: {template_path}")