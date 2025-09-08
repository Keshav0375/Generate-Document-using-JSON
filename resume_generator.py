"""
Resume Generator - Main Module
Converts JSON data to formatted Word and PDF documents
"""

import json
import os
from typing import Dict, List, Optional, Tuple
from datetime import datetime
from pathlib import Path
import logging

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx2pdf
from dataclasses import dataclass, field
from typing import Union

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class DocumentConfig:
    """Configuration for document formatting"""
    # Page margins (in inches)
    margin_top: float = 0.5
    margin_bottom: float = 0.5
    margin_left: float = 0.5
    margin_right: float = 0.5

    # Font settings
    font_name: str = "Calibri"
    font_size_normal: int = 11
    font_size_heading: int = 14
    font_size_section: int = 12
    font_size_name: int = 20

    # Spacing
    line_spacing: float = 1.0
    paragraph_spacing_before: int = 0
    paragraph_spacing_after: int = 0
    section_spacing: int = 6

    # Colors
    heading_color: Tuple[int, int, int] = (0, 0, 0)
    text_color: Tuple[int, int, int] = (0, 0, 0)
    link_color: Tuple[int, int, int] = (0, 0, 139)


class ResumeFormatter:
    """Handles formatting operations for the resume document"""

    def __init__(self, config: DocumentConfig = None):
        self.config = config or DocumentConfig()

    def set_margins(self, doc: Document) -> None:
        """Set document margins"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(self.config.margin_top)
            section.bottom_margin = Inches(self.config.margin_bottom)
            section.left_margin = Inches(self.config.margin_left)
            section.right_margin = Inches(self.config.margin_right)

    def add_horizontal_line(self, doc: Document) -> None:
        """Add a horizontal line to the document"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pBorder = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBorder.append(bottom)
        p._p.get_or_add_pPr().append(pBorder)

    def format_paragraph(self, paragraph, font_size: int = None,
                         bold: bool = False, italic: bool = False,
                         color: Tuple[int, int, int] = None) -> None:
        """Format a paragraph with specified styles"""
        if paragraph.runs:
            for run in paragraph.runs:
                run.font.name = self.config.font_name
                run.font.size = Pt(font_size or self.config.font_size_normal)
                run.font.bold = bold
                run.font.italic = italic
                if color:
                    run.font.color.rgb = RGBColor(*color)

    def set_paragraph_spacing(self, paragraph, before: int = None, after: int = None) -> None:
        """Set paragraph spacing"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(before or self.config.paragraph_spacing_before)
        paragraph_format.space_after = Pt(after or self.config.paragraph_spacing_after)
        paragraph_format.line_spacing = self.config.line_spacing


class ResumeBuilder:
    """Main class for building resume documents"""

    def __init__(self, config: DocumentConfig = None):
        self.config = config or DocumentConfig()
        self.formatter = ResumeFormatter(self.config)
        self.doc = None

    def create_document(self) -> Document:
        """Create a new document with basic setup"""
        self.doc = Document()
        self.formatter.set_margins(self.doc)
        return self.doc

    def add_header(self, data: Dict) -> None:
        """Add header section with name and contact info"""
        if not self.doc:
            raise ValueError("Document not created. Call create_document() first.")

        # Add name
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(data.get('name', ''))
        name_run.font.name = self.config.font_name
        name_run.font.size = Pt(self.config.font_size_name)
        name_run.font.bold = True

        # Add contact info line
        contact_parts = []
        if data.get('phone'):
            contact_parts.append(data['phone'])
        if data.get('email'):
            contact_parts.append(data['email'])
        if data.get('location'):
            contact_parts.append(data['location'])

        if contact_parts:
            contact_para = self.doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(' | '.join(contact_parts))
            contact_run.font.name = self.config.font_name
            contact_run.font.size = Pt(self.config.font_size_normal)

        # Add links line
        links = []
        if data.get('linkedin'):
            links.append(f"LinkedIn")
        if data.get('portfolio'):
            links.append(f"Portfolio")
        if data.get('github'):
            links.append(f"GitHub")

        if links:
            links_para = self.doc.add_paragraph()
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links_run = links_para.add_run(' | '.join(links))
            links_run.font.name = self.config.font_name
            links_run.font.size = Pt(self.config.font_size_normal)

        # Add spacing after header
        self.doc.add_paragraph()

    def add_section_header(self, title: str) -> None:
        """Add a section header with formatting"""
        header_para = self.doc.add_paragraph()
        header_run = header_para.add_run(title.upper())
        header_run.font.name = self.config.font_name
        header_run.font.size = Pt(self.config.font_size_section)
        header_run.font.bold = True

        # Add underline
        self.formatter.add_horizontal_line(self.doc)

    def add_technical_skills(self, skills: Dict[str, str]) -> None:
        """Add technical skills section"""
        self.add_section_header("Technical Skills")

        for category, skill_list in skills.items():
            para = self.doc.add_paragraph()

            # Add category in bold
            category_run = para.add_run(f"{category:<15}")
            category_run.font.name = self.config.font_name
            category_run.font.size = Pt(self.config.font_size_normal)
            category_run.font.bold = True

            # Add skills
            skills_run = para.add_run(skill_list)
            skills_run.font.name = self.config.font_name
            skills_run.font.size = Pt(self.config.font_size_normal)

    def add_education(self, education: List[Dict]) -> None:
        """Add education section"""
        self.add_section_header("Education")

        for edu in education:
            # Degree and dates on same line
            para = self.doc.add_paragraph()
            degree_run = para.add_run(edu.get('degree', ''))
            degree_run.font.name = self.config.font_name
            degree_run.font.size = Pt(self.config.font_size_normal)
            degree_run.font.bold = True

            # Add tabs for alignment
            para.add_run('\t' * 2)
            date_run = para.add_run(edu.get('dates', ''))
            date_run.font.name = self.config.font_name
            date_run.font.size = Pt(self.config.font_size_normal)

            # School and location
            school_para = self.doc.add_paragraph()
            school_text = f"{edu.get('school', '')} — {edu.get('gpa', '')}" if edu.get('gpa') else edu.get('school', '')
            school_run = school_para.add_run(school_text)
            school_run.font.name = self.config.font_name
            school_run.font.size = Pt(self.config.font_size_normal)

            school_para.add_run('\t' * 3)
            location_run = school_para.add_run(edu.get('location', ''))
            location_run.font.name = self.config.font_name
            location_run.font.size = Pt(self.config.font_size_normal)

            # Add additional notes as bullet points
            if edu.get('notes'):
                for note in edu['notes']:
                    bullet_para = self.doc.add_paragraph(style='List Bullet')
                    bullet_run = bullet_para.add_run(note)
                    bullet_run.font.name = self.config.font_name
                    bullet_run.font.size = Pt(self.config.font_size_normal)

    def add_experience(self, experiences: List[Dict]) -> None:
        """Add experience section"""
        self.add_section_header("Experience")

        for exp in experiences:
            # Job title
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run(exp.get('title', ''))
            title_run.font.name = self.config.font_name
            title_run.font.size = Pt(self.config.font_size_normal)
            title_run.font.bold = True

            # Add dates aligned to right
            title_para.add_run('\t' * 3)
            date_run = title_para.add_run(exp.get('dates', ''))
            date_run.font.name = self.config.font_name
            date_run.font.size = Pt(self.config.font_size_normal)

            # Company and location
            company_para = self.doc.add_paragraph()
            company_run = company_para.add_run(exp.get('company', ''))
            company_run.font.name = self.config.font_name
            company_run.font.size = Pt(self.config.font_size_normal)
            company_run.font.italic = True

            company_para.add_run('\t' * 4)
            location_run = company_para.add_run(exp.get('location', ''))
            location_run.font.name = self.config.font_name
            location_run.font.size = Pt(self.config.font_size_normal)
            location_run.font.italic = True

            # Add bullet points
            for bullet in exp.get('bullets', []):
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_run = bullet_para.add_run(bullet)
                bullet_run.font.name = self.config.font_name
                bullet_run.font.size = Pt(self.config.font_size_normal)

            # Add spacing between experiences
            if experiences.index(exp) < len(experiences) - 1:
                self.doc.add_paragraph()

    def add_projects(self, projects: List[Dict]) -> None:
        """Add projects section"""
        self.add_section_header("Projects")

        for project in projects:
            # Project name and dates
            proj_para = self.doc.add_paragraph()

            # Project name and description
            name_text = project.get('name', '')
            if project.get('description'):
                name_text += f" | {project['description']}"

            proj_run = proj_para.add_run(name_text)
            proj_run.font.name = self.config.font_name
            proj_run.font.size = Pt(self.config.font_size_normal)
            proj_run.font.bold = True

            # Dates
            proj_para.add_run('\t')
            date_run = proj_para.add_run(project.get('dates', ''))
            date_run.font.name = self.config.font_name
            date_run.font.size = Pt(self.config.font_size_normal)

            # Bullet points
            for bullet in project.get('bullets', []):
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_run = bullet_para.add_run(bullet)
                bullet_run.font.name = self.config.font_name
                bullet_run.font.size = Pt(self.config.font_size_normal)

    def add_competitions(self, competitions: List[Dict]) -> None:
        """Add coding competitions section"""
        self.add_section_header("Coding Competitions")

        for comp in competitions:
            # Competition name and date
            comp_para = self.doc.add_paragraph()
            comp_run = comp_para.add_run(comp.get('name', ''))
            comp_run.font.name = self.config.font_name
            comp_run.font.size = Pt(self.config.font_size_normal)
            comp_run.font.bold = True

            comp_para.add_run('\t' * 4)
            date_run = comp_para.add_run(comp.get('date', ''))
            date_run.font.name = self.config.font_name
            date_run.font.size = Pt(self.config.font_size_normal)

            # Organization and location
            org_para = self.doc.add_paragraph()
            org_run = org_para.add_run(comp.get('organization', ''))
            org_run.font.name = self.config.font_name
            org_run.font.size = Pt(self.config.font_size_normal)

            org_para.add_run('\t' * 4)
            loc_run = org_para.add_run(comp.get('location', ''))
            loc_run.font.name = self.config.font_name
            loc_run.font.size = Pt(self.config.font_size_normal)

            # Description bullets
            for bullet in comp.get('bullets', []):
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_run = bullet_para.add_run(bullet)
                bullet_run.font.name = self.config.font_name
                bullet_run.font.size = Pt(self.config.font_size_normal)

    def add_certifications(self, certifications: List[Dict]) -> None:
        """Add certifications section"""
        self.add_section_header("Certifications")

        for cert in certifications:
            cert_para = self.doc.add_paragraph(style='List Bullet')
            cert_text = cert.get('name', '')

            cert_run = cert_para.add_run(cert_text)
            cert_run.font.name = self.config.font_name
            cert_run.font.size = Pt(self.config.font_size_normal)

            # Add date aligned to right
            cert_para.add_run('\t' * 4)
            date_run = cert_para.add_run(cert.get('date', ''))
            date_run.font.name = self.config.font_name
            date_run.font.size = Pt(self.config.font_size_normal)

    def build_resume(self, resume_data: Dict) -> Document:
        """Build complete resume from JSON data"""
        try:
            self.create_document()

            # Add sections in order
            if 'header' in resume_data:
                self.add_header(resume_data['header'])

            if 'technical_skills' in resume_data:
                self.add_technical_skills(resume_data['technical_skills'])
                self.doc.add_paragraph()  # Add spacing

            if 'education' in resume_data:
                self.add_education(resume_data['education'])
                self.doc.add_paragraph()

            if 'experience' in resume_data:
                self.add_experience(resume_data['experience'])
                self.doc.add_paragraph()

            if 'projects' in resume_data:
                self.add_projects(resume_data['projects'])
                self.doc.add_paragraph()

            if 'competitions' in resume_data:
                self.add_competitions(resume_data['competitions'])
                self.doc.add_paragraph()

            if 'certifications' in resume_data:
                self.add_certifications(resume_data['certifications'])

            return self.doc

        except Exception as e:
            logger.error(f"Error building resume: {str(e)}")
            raise


class ResumeGenerator:
    """Main class for generating resume files"""

    def __init__(self, config: DocumentConfig = None):
        self.config = config or DocumentConfig()
        self.builder = ResumeBuilder(self.config)

    def load_json(self, json_path: str) -> Dict:
        """Load resume data from JSON file"""
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            logger.error(f"JSON file not found: {json_path}")
            raise
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON format: {str(e)}")
            raise

    def save_json(self, data: Dict, json_path: str) -> None:
        """Save resume data to JSON file"""
        try:
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            logger.info(f"JSON saved to {json_path}")
        except Exception as e:
            logger.error(f"Error saving JSON: {str(e)}")
            raise

    def generate_word(self, resume_data: Dict, output_path: str) -> str:
        """Generate Word document from resume data"""
        try:
            doc = self.builder.build_resume(resume_data)
            doc.save(output_path)
            logger.info(f"Word document saved to {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            raise

    def generate_pdf(self, word_path: str, pdf_path: str = None) -> str:
        """Convert Word document to PDF"""
        try:
            if pdf_path is None:
                pdf_path = word_path.replace('.docx', '.pdf')

            # Convert to PDF
            docx2pdf.convert(word_path, pdf_path)
            logger.info(f"PDF saved to {pdf_path}")
            return pdf_path
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            # If docx2pdf fails, return None but don't crash
            logger.warning("PDF generation failed. Word document is still available.")
            return None

    def generate_from_json(self, json_path: str, output_dir: str = None,
                           base_name: str = "resume") -> Tuple[str, Optional[str]]:
        """Generate both Word and PDF from JSON file"""
        try:
            # Load JSON data
            resume_data = self.load_json(json_path)

            # Set output directory
            if output_dir is None:
                output_dir = os.path.dirname(json_path) or '.'

            # Create output directory if it doesn't exist
            Path(output_dir).mkdir(parents=True, exist_ok=True)

            # Generate file paths
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            word_path = os.path.join(output_dir, f"{base_name}_{timestamp}.docx")
            pdf_path = os.path.join(output_dir, f"{base_name}_{timestamp}.pdf")

            # Generate Word document
            self.generate_word(resume_data, word_path)

            # Try to generate PDF
            pdf_result = self.generate_pdf(word_path, pdf_path)

            return word_path, pdf_result

        except Exception as e:
            logger.error(f"Error in generate_from_json: {str(e)}")
            raise


# Example usage
if __name__ == "__main__":
    # Create generator with custom config
    config = DocumentConfig(
        margin_top=0.5,
        margin_bottom=0.5,
        margin_left=0.5,
        margin_right=0.5,
        font_name="Calibri",
        font_size_normal=11
    )

    generator = ResumeGenerator(config)

    # Generate resume
    try:
        word_file, pdf_file = generator.generate_from_json(
            json_path="resume_data.json",
            output_dir="output",
            base_name="keshav_arri_resume"
        )
        print(f"✅ Word document: {word_file}")
        if pdf_file:
            print(f"✅ PDF document: {pdf_file}")
        else:
            print("⚠️ PDF generation failed, but Word document is available")
    except Exception as e:
        print(f"❌ Error: {str(e)}")